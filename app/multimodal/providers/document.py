"""TXT、Markdown、DOCX 与 PDF 文字层的本地异步解析 Provider。"""

from __future__ import annotations

import asyncio
from pathlib import PurePosixPath

from docx import Document
from openpyxl import load_workbook  # type: ignore[import-untyped]
from pypdf import PdfReader

from app.multimodal.errors import MaterialIngestError
from app.multimodal.models import DownloadedFile, MaterialLocator, ProviderResult, ProviderSegment
from app.multimodal.providers.base import DocumentParser


class LocalDocumentParser(DocumentParser):
    def __init__(self, *, max_document_chars: int, minimum_pdf_text_chars: int = 20) -> None:
        self.max_document_chars = max_document_chars
        self.minimum_pdf_text_chars = minimum_pdf_text_chars

    async def parse(self, source: DownloadedFile) -> ProviderResult:
        return await asyncio.to_thread(self._parse_sync, source)

    def _parse_sync(self, source: DownloadedFile) -> ProviderResult:
        suffix = PurePosixPath(source.filename).suffix.lower()
        try:
            if suffix in {".txt", ".md"}:
                return self._parse_text(source)
            if suffix == ".docx":
                return self._parse_docx(source)
            if suffix == ".xlsx":
                return self._parse_xlsx(source)
            if suffix == ".pdf":
                return self._parse_pdf(source)
        except MaterialIngestError:
            raise
        except Exception as exc:
            raise MaterialIngestError(
                "XDW-DOC-PARSE", "文档无法解析，请检查文件是否损坏。"
            ) from exc
        raise MaterialIngestError("XDW-DOC-TYPE", "当前文档类型不受支持。")

    def _parse_text(self, source: DownloadedFile) -> ProviderResult:
        raw = source.path.read_bytes()
        text: str | None = None
        encodings = ["utf-8-sig"]
        if raw.startswith((b"\xff\xfe", b"\xfe\xff")):
            encodings.append("utf-16")
        encodings.append("gb18030")
        for encoding in encodings:
            try:
                text = raw.decode(encoding)
                break
            except UnicodeDecodeError:
                continue
        if text is None:
            raise MaterialIngestError("XDW-DOC-ENCODING", "文本编码无法识别，请转为 UTF-8。")
        normalized = _normalize_text(text)
        self._validate_text(normalized)
        return ProviderResult(
            provider_name="local",
            provider_model="text-parser-v1",
            normalized_text=normalized,
            segments=_text_segments(normalized),
        )

    def _parse_docx(self, source: DownloadedFile) -> ProviderResult:
        document = Document(str(source.path))
        blocks: list[str] = [paragraph.text.strip() for paragraph in document.paragraphs]
        for table in document.tables:
            blocks.extend(
                "｜".join(cell.text.strip() for cell in row.cells).strip("｜")
                for row in table.rows
            )
        normalized = "\n\n".join(block for block in blocks if block)
        normalized = _normalize_text(normalized)
        self._validate_text(normalized)
        return ProviderResult(
            provider_name="local",
            provider_model="docx-parser-v1",
            normalized_text=normalized,
            segments=_block_segments(normalized),
        )

    def _parse_xlsx(self, source: DownloadedFile) -> ProviderResult:
        """将表格转为带工作表、行号的可引用文本，而不是让模型猜测表格结构。"""

        workbook = load_workbook(filename=source.path, read_only=True, data_only=True)
        blocks: list[str] = []
        try:
            for worksheet in workbook.worksheets[:20]:
                rows: list[str] = []
                for row_number, values in enumerate(worksheet.iter_rows(values_only=True), start=1):
                    rendered = [str(value).strip() for value in values if value is not None and str(value).strip()]
                    if rendered:
                        rows.append(f"第 {row_number} 行：" + " | ".join(rendered))
                    if row_number >= 10_000:
                        raise MaterialIngestError(
                            "XDW-XLSX-ROWS", "单个工作表超过 10000 行，请拆分后再上传。"
                        )
                if rows:
                    blocks.append(f"[工作表：{worksheet.title}]\n" + "\n".join(rows))
        finally:
            workbook.close()

        normalized = _normalize_text("\n\n".join(blocks))
        self._validate_text(normalized)
        return ProviderResult(
            provider_name="local",
            provider_model="xlsx-parser-v1",
            normalized_text=normalized,
            segments=_block_segments(normalized),
        )

    def _parse_pdf(self, source: DownloadedFile) -> ProviderResult:
        reader = PdfReader(source.path)
        pages = [_normalize_text(page.extract_text() or "") for page in reader.pages]
        visible_pages = [(index, text) for index, text in enumerate(pages, start=1) if text]
        normalized = "\n\n".join(text for _, text in visible_pages)
        if len(normalized) < self.minimum_pdf_text_chars:
            raise MaterialIngestError(
                "XDW-DOC-OCR-REQUIRED",
                "PDF 没有足够的文字层，需要在 OCR 阶段处理。",
            )
        self._validate_text(normalized)
        segments: list[ProviderSegment] = []
        offset = 0
        for page, text in visible_pages:
            for start, end, chunk in _chunks(text):
                segments.append(
                    ProviderSegment(
                        text=chunk,
                        confidence=1.0,
                        locator=MaterialLocator(
                            page=page,
                            char_start=offset + start,
                            char_end=offset + end,
                        ),
                    )
                )
            offset += len(text) + 2
        return ProviderResult(
            provider_name="local",
            provider_model="pdf-text-parser-v1",
            normalized_text=normalized,
            segments=segments,
        )

    def _validate_text(self, text: str) -> None:
        if not text:
            raise MaterialIngestError("XDW-DOC-EMPTY", "文档没有可解析的正文。")
        if len(text) > self.max_document_chars:
            raise MaterialIngestError("XDW-DOC-CHARS", "文档正文超过字符数量限制。")


def _normalize_text(text: str) -> str:
    return text.replace("\r\n", "\n").replace("\r", "\n").replace("\x00", "").strip()


def _chunks(text: str, size: int = 50_000) -> list[tuple[int, int, str]]:
    return [(start, min(start + size, len(text)), text[start : start + size]) for start in range(0, len(text), size)]


def _text_segments(text: str) -> list[ProviderSegment]:
    return [
        ProviderSegment(
            text=chunk,
            confidence=1.0,
            locator=MaterialLocator(char_start=start, char_end=end),
        )
        for start, end, chunk in _chunks(text)
    ]


def _block_segments(text: str) -> list[ProviderSegment]:
    segments: list[ProviderSegment] = []
    cursor = 0
    for block in text.split("\n\n"):
        start = text.find(block, cursor)
        end = start + len(block)
        for chunk_start, chunk_end, chunk in _chunks(block):
            segments.append(
                ProviderSegment(
                    text=chunk,
                    confidence=1.0,
                    locator=MaterialLocator(
                        char_start=start + chunk_start,
                        char_end=start + chunk_end,
                    ),
                )
            )
        cursor = end
    return segments
