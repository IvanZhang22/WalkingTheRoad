from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader


class DocumentParseError(ValueError):
    pass


SUPPORTED_EXTENSIONS = {".txt", ".md", ".docx", ".pdf"}


def parse_document(
    filename: str,
    content: bytes,
    *,
    max_upload_bytes: int,
    max_document_chars: int,
) -> str:
    safe_name = Path(filename or "").name
    extension = Path(safe_name).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise DocumentParseError("仅支持 TXT、MD、DOCX 和带文字层的 PDF 文件。")
    if not content:
        raise DocumentParseError("上传文件为空。")
    if len(content) > max_upload_bytes:
        raise DocumentParseError(
            f"文件超过 {max_upload_bytes // 1024 // 1024} MB 限制，请压缩或拆分材料。"
        )

    try:
        if extension in {".txt", ".md"}:
            text = _decode_text(content)
        elif extension == ".docx":
            text = _read_docx(content)
        else:
            text = _read_pdf(content)
    except DocumentParseError:
        raise
    except Exception as exc:
        raise DocumentParseError(f"文件解析失败：{exc}") from exc

    text = text.replace("\x00", "").strip()
    if not text:
        if extension == ".pdf":
            raise DocumentParseError("PDF 没有可提取的文字层；扫描件需在后续 OCR 版本处理。")
        raise DocumentParseError("文件中没有提取到有效正文。")
    if len(text) > max_document_chars:
        raise DocumentParseError(
            f"正文共 {len(text)} 字符，超过当前 {max_document_chars} 字符限制，请拆分材料。"
        )
    return text


def _decode_text(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentParseError("文本文件编码无法识别，请转换为 UTF-8。")


def _read_docx(content: bytes) -> str:
    document = Document(BytesIO(content))
    chunks = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            chunks.append("｜".join(cell.text.strip() for cell in row.cells))
    return "\n".join(chunks)


def _read_pdf(content: bytes) -> str:
    reader = PdfReader(BytesIO(content))
    return "\n\n".join((page.extract_text() or "").strip() for page in reader.pages)
