from io import BytesIO

import pytest
from docx import Document
from pypdf import PdfWriter

from app.document_parser import DocumentParseError, parse_document

LIMITS = {"max_upload_bytes": 20 * 1024 * 1024, "max_document_chars": 300_000}


def test_txt_and_markdown_decoding() -> None:
    assert parse_document("材料.txt", "访谈正文".encode(), **LIMITS) == "访谈正文"
    assert parse_document("材料.md", "# I01\n正文".encode(), **LIMITS).startswith("# I01")


def test_docx_paragraphs_and_tables() -> None:
    document = Document()
    document.add_paragraph("I01｜访谈材料")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "字段"
    table.cell(0, 1).text = "内容"
    stream = BytesIO()
    document.save(stream)
    text = parse_document("材料.docx", stream.getvalue(), **LIMITS)
    assert "I01｜访谈材料" in text
    assert "字段｜内容" in text


def test_blank_pdf_reports_ocr_boundary() -> None:
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    stream = BytesIO()
    writer.write(stream)
    with pytest.raises(DocumentParseError, match="文字层"):
        parse_document("扫描件.pdf", stream.getvalue(), **LIMITS)


@pytest.mark.parametrize("name", ["材料.png", "材料.wav", "材料.exe"])
def test_unsupported_extension(name: str) -> None:
    with pytest.raises(DocumentParseError, match="仅支持"):
        parse_document(name, b"123", **LIMITS)


def test_size_and_character_limits() -> None:
    with pytest.raises(DocumentParseError, match="文件超过"):
        parse_document("材料.txt", b"12345", max_upload_bytes=4, max_document_chars=100)
    with pytest.raises(DocumentParseError, match="正文共"):
        parse_document("材料.txt", b"12345", max_upload_bytes=10, max_document_chars=4)
