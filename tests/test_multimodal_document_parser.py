from __future__ import annotations

import hashlib
from pathlib import Path

import httpx
import pytest
from docx import Document
from openpyxl import Workbook
from pypdf import PdfWriter

from app.multimodal.contracts import FileContentPart
from app.multimodal.downloader import SafeDownloader
from app.multimodal.errors import MaterialIngestError
from app.multimodal.models import DownloadedFile, MaterialModality, MaterialStatus
from app.multimodal.providers.document import LocalDocumentParser
from app.multimodal.providers.mock import MockASRProvider, MockOCRProvider
from app.multimodal.service import MaterialIngestService


def downloaded(path: Path, filename: str, mime_type: str) -> DownloadedFile:
    content = path.read_bytes()
    return DownloadedFile(
        path=path,
        filename=filename,
        mime_type=mime_type,
        size_bytes=len(content),
        sha256=hashlib.sha256(content).hexdigest(),
    )


async def test_parses_utf8_gb18030_and_preserves_character_locations(tmp_path: Path) -> None:
    parser = LocalDocumentParser(max_document_chars=1000)
    utf8_path = tmp_path / "utf8.txt"
    utf8_path.write_bytes("第一行\r\n第二行".encode())
    utf8 = await parser.parse(downloaded(utf8_path, "utf8.txt", "text/plain"))
    assert utf8.normalized_text == "第一行\n第二行"
    assert utf8.segments[0].locator.char_start == 0
    assert utf8.segments[0].locator.char_end == len(utf8.normalized_text)

    gb_path = tmp_path / "gb.txt"
    gb_path.write_bytes("中文编码".encode("gb18030"))
    gb = await parser.parse(downloaded(gb_path, "gb.txt", "text/plain"))
    assert gb.normalized_text == "中文编码"


async def test_parses_docx_paragraphs_and_tables(tmp_path: Path) -> None:
    path = tmp_path / "材料.docx"
    document = Document()
    document.add_paragraph("第一段")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "字段"
    table.cell(0, 1).text = "内容"
    document.save(path)
    result = await LocalDocumentParser(max_document_chars=1000).parse(
        downloaded(
            path,
            "材料.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    )
    assert result.normalized_text == "第一段\n\n字段｜内容"
    assert len(result.segments) == 2
    assert result.segments[1].locator.char_start == len("第一段\n\n")


async def test_parses_xlsx_with_sheet_and_row_locations(tmp_path: Path) -> None:
    path = tmp_path / "material.xlsx"
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "interview"
    worksheet.append(["participant", "statement"])
    worksheet.append(["N01", "Service entry is difficult to find"])
    workbook.save(path)
    result = await LocalDocumentParser(max_document_chars=1000).parse(
        downloaded(
            path,
            "material.xlsx",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
    )
    assert "[工作表：interview]" in result.normalized_text
    assert "第 2 行：N01 | Service entry is difficult to find" in result.normalized_text
    assert result.segments[0].locator.char_start == 0


async def test_pdf_pages_keep_page_and_global_character_locations(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    class Page:
        def __init__(self, text: str) -> None:
            self.text = text

        def extract_text(self) -> str:
            return self.text

    monkeypatch.setattr(
        "app.multimodal.providers.document.PdfReader",
        lambda path: type("Reader", (), {"pages": [Page("第一页正文足够长"), Page("第二页正文足够长")]})(),
    )
    path = tmp_path / "材料.pdf"
    path.write_bytes(b"%PDF-1.7")
    result = await LocalDocumentParser(
        max_document_chars=1000, minimum_pdf_text_chars=10
    ).parse(downloaded(path, "材料.pdf", "application/pdf"))
    assert [segment.locator.page for segment in result.segments] == [1, 2]
    assert result.segments[1].locator.char_start == len("第一页正文足够长\n\n")


async def test_blank_pdf_routes_to_future_ocr_and_character_limit_is_enforced(
    tmp_path: Path,
) -> None:
    path = tmp_path / "扫描件.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    writer.write(path)
    parser = LocalDocumentParser(max_document_chars=5)
    with pytest.raises(MaterialIngestError) as scan:
        await parser.parse(downloaded(path, "扫描件.pdf", "application/pdf"))
    assert scan.value.code == "XDW-DOC-OCR-REQUIRED"

    text_path = tmp_path / "long.txt"
    text_path.write_text("超过五个字符", encoding="utf-8")
    with pytest.raises(MaterialIngestError) as long_text:
        await parser.parse(downloaded(text_path, "long.txt", "text/plain"))
    assert long_text.value.code == "XDW-DOC-CHARS"


async def test_real_document_service_cleans_temp_file_and_returns_material(tmp_path: Path) -> None:
    content = "学生认为信息入口分散。".encode()
    transport = httpx.MockTransport(
        lambda request: httpx.Response(200, headers={"content-type": "text/plain"}, content=content)
    )

    async def resolver(hostname: str, port: int) -> list[str]:
        return ["93.184.216.34"]

    service = MaterialIngestService(
        downloader=SafeDownloader(
            max_bytes=1024,
            tmp_dir=tmp_path,
            resolver=resolver,
            transport=transport,
            retries=0,
        ),
        asr=MockASRProvider(),
        ocr=MockOCRProvider(),
        document_parser=LocalDocumentParser(max_document_chars=1000),
        enabled_modalities=frozenset({MaterialModality.document}),
    )
    part = FileContentPart.model_validate(
        {
            "type": "file",
            "file": {"url": "https://files.example.org/a.txt", "filename": "a.txt"},
        }
    )
    material = (await service.ingest([part]))[0]
    assert material.status is MaterialStatus.ready
    assert material.normalized_text == "学生认为信息入口分散。"
    assert material.automatic_text == material.normalized_text
    assert material.segments[0].locator.char_end == len(material.normalized_text)
    assert list(tmp_path.iterdir()) == []


async def test_parser_failure_also_cleans_downloaded_temp_file(tmp_path: Path) -> None:
    transport = httpx.MockTransport(
        lambda request: httpx.Response(
            200, headers={"content-type": "text/plain"}, content=b""
        )
    )

    async def resolver(hostname: str, port: int) -> list[str]:
        return ["93.184.216.34"]

    service = MaterialIngestService(
        downloader=SafeDownloader(
            max_bytes=1024,
            tmp_dir=tmp_path,
            resolver=resolver,
            transport=transport,
            retries=0,
        ),
        asr=MockASRProvider(),
        ocr=MockOCRProvider(),
        document_parser=LocalDocumentParser(max_document_chars=1000),
        enabled_modalities=frozenset({MaterialModality.document}),
    )
    part = FileContentPart.model_validate(
        {
            "type": "file",
            "file": {"url": "https://files.example.org/empty.txt", "filename": "empty.txt"},
        }
    )
    material = (await service.ingest([part]))[0]
    assert material.status is MaterialStatus.failed
    assert material.issues[0].code == "XDW-DOC-EMPTY"
    assert list(tmp_path.iterdir()) == []
